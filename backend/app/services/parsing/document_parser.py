"""File → raw text extraction for PDF, DOCX and plain text."""
from __future__ import annotations

import io
from dataclasses import dataclass

from app.core.exceptions import DocumentParseError, EmptyDocumentError, UnsupportedFileType


@dataclass
class ParsedDocument:
    text: str
    page_count: int


def parse_document(filename: str, content: bytes) -> ParsedDocument:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        doc = _parse_pdf(content)
    elif name.endswith(".docx"):
        doc = _parse_docx(content)
    elif name.endswith((".txt", ".md")):
        doc = ParsedDocument(text=content.decode("utf-8", errors="replace"), page_count=1)
    else:
        raise UnsupportedFileType(
            f"Unsupported file type for '{filename}'. Upload a PDF, DOCX or TXT file."
        )

    if not doc.text.strip():
        raise EmptyDocumentError(
            f"No text could be extracted from '{filename}'. "
            "If this is a scanned/image-only PDF, export a text-based version."
        )
    return doc


def _parse_pdf(content: bytes) -> ParsedDocument:
    # PyMuPDF is fast and layout-robust; pdfplumber is the fallback for PDFs
    # PyMuPDF struggles with (some generator quirks, odd encodings).
    try:
        import fitz  # PyMuPDF

        with fitz.open(stream=content, filetype="pdf") as pdf:
            pages = [page.get_text("text") for page in pdf]
        text = "\n".join(pages)
        if text.strip():
            return ParsedDocument(text=text, page_count=len(pages))
    except Exception:
        pass

    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        return ParsedDocument(text="\n".join(pages), page_count=len(pages))
    except Exception as exc:
        raise DocumentParseError(f"Failed to read PDF: {exc}") from exc


def _parse_docx(content: bytes) -> ParsedDocument:
    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        parts: list[str] = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        text = "\n".join(parts)
        # DOCX has no fixed pages; estimate for the length heuristics.
        page_count = max(1, round(len(text.split()) / 500))
        return ParsedDocument(text=text, page_count=page_count)
    except Exception as exc:
        raise DocumentParseError(f"Failed to read DOCX: {exc}") from exc
